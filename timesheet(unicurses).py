import csv,re,xlsxwriter
from docx import *
from datetime import datetime, timedelta
from unicurses import *


Input, Data, Date, Time, Door, Event, Zone = [],[],[],[],[],[],[]
mon, tue, wed, thu, fri, sat, sun = [False],[False],[False],[False],[False],[True],[True]
week = ['Monday','Tuesday','Wednesday', 'Thursday','Friday','Saturday','Sunday']


def overtime(Day,In,Out):
    

    #print delta
    #print Day
##    print h
##    print m
##    print s
    if Day != 'Friday' and Day != 'Saturday' and Day != 'Sunday':
        print Day
        In = '16:30:00'
        t1 = datetime.strptime(In,'%H:%M:%S')
        t2 = datetime.strptime(Out,'%H:%M:%S')

        delta = (t2 - t1)
        
        print t1
        print t2
        print delta
        s = delta.seconds
        h, s = divmod(s, 3600)
        m, s = divmod(s, 60)

        if m >= 30:
            return float(h)+0.5
        else:
            return h
        
    elif Day == 'Friday':
        print 'Friday'
        In = '16:00:00'
        
        t1 = datetime.strptime(In,'%H:%M:%S')
        t2 = datetime.strptime(Out,'%H:%M:%S')
        print t1
        delta = (t2 - t1)
        print delta
        s = delta.seconds
        h, s = divmod(s, 3600)
        m, s = divmod(s, 60)

        if m >= 30:
            return float(h)+0.5
        else:
            return h

    else:
        t1 = datetime.strptime(In,'%H:%M:%S')
        t2 = datetime.strptime(Out,'%H:%M:%S')

        delta = (t2 - t1)

        s = delta.seconds
        h, s = divmod(s, 3600)
        m, s = divmod(s, 60)
        
        
        if m >= 30:
            return float(h)+0.5
        else:
            return h





with open('filename2.csv') as fp:
    for i, line in enumerate(fp):
        if i == 2:
            emp = line
            break
searchPattern = re.compile('Report On:(.*?)\(')
for i in re.findall(searchPattern,str(emp)):
    Name = i.strip()
searchPattern = re.compile('\((.*?)\)')
for i in re.findall(searchPattern,str(emp)):
    Num = i.strip()
  


with open('filename2.csv','r') as f:
    f.next()
    f.next()
    f.next()
    f.next()
    f.next() # skip 5 lines
    for line in f:
        Input.append(line)


read = csv.reader(Input)
for row in read:
    Date.append(row[0])
    Time.append(row[1])
    Door.append(row[2])
    Event.append(row[3])
    Zone.append(row[5])

prev = Date[0]
Data.append(Date[0])
Data.append(Time[0])

for index in range(len(Date)): 
    if Date[index]!= prev:
        Data.append(Time[index-1])
        Data.append(Date[index])
        Data.append(Time[index])
        prev = Date[index]
Data.append(Time[len(Time)-1])


for index in range(len(Data)):
    if Data[index][:3] == 'Mon':
        mon.append(Data[index][4:])
        mon.append(Data[index+1])
        mon.append(Data[index+2])
    elif Data[index][:3] == 'Tue':
        tue.append(Data[index][4:])
        tue.append(Data[index+1])
        tue.append(Data[index+2])
    elif Data[index][:3] == 'Wed':
        wed.append(Data[index][4:])
        wed.append(Data[index+1])
        wed.append(Data[index+2])
    elif Data[index][:3] == 'Thu':
        thu.append(Data[index][4:])
        thu.append(Data[index+1])
        thu.append(Data[index+2])
    elif Data[index][:3] == 'Fri':
        fri.append(Data[index][4:])
        fri.append(Data[index+1])
        fri.append(Data[index+2])
    elif Data[index][:3] == 'Sat':
        sat.append(Data[index][4:])
        sat.append(Data[index+1])
        sat.append(Data[index+2])
    elif Data[index][:3] == 'Sun':
        sun.append(Data[index][4:])
        sun.append(Data[index+1])
        sun.append(Data[index+2])

#start = raw_input("Please enter week start date (ddmmyy) :")
##end = raw_input("Please enter week end date   (ddmmyy) :")
##start = datetime.strptime(end, '%d%m%y') - timedelta(days=6)
##start = start.strftime('%Y-%m-%d')

########UNICURSES TEST#############
stdscr = initscr()
attron(A_BOLD)
addstr('Processing Time and Attendance details of '+Name+' - '+Num+'\n\n\n')
attroff(A_BOLD)
##addstr('Please enter week end date   (ddmmyy) :')
##end = getstr()
##start = datetime.strptime(end, '%d%m%y') - timedelta(days=6)
##start = start.strftime('%Y-%m-%d')


while True:
    c = getch()
    print c
##    if c == ord('q'):
##        break  # Exit the while loop
    if c == KEY_UP:
        addstr('up key')
        break
    elif c == KEY_DOWN:
        print "down key"


addstr(end)
addstr(start)
a = getch()


 # Create a workbook and add a worksheet.
workbook = xlsxwriter.Workbook('OT-'+Name+'-'+str(datetime.now().strftime("%d-%m-%Y"))+'.xlsx')
sheet1 = workbook.add_worksheet('HR Sheet')
sheet2 = workbook.add_worksheet('T&A')

 # Add  formats to use to highlight cells.
headformat = workbook.add_format({'bold': True, 'italic': True, 'border': 2,'font_size': 16,'text_wrap': True,'align':'center'})
fieldformat = workbook.add_format({'left': 2,'right': 2,'top':1,'bottom':1,'text_wrap': True,'align':'center'})
borbol = workbook.add_format({'bold': True, 'border': True})


 # Adjust the column width.
                #( , ,width)
sheet1.set_column(0, 0, 11.57)# Emp no.
sheet1.set_column(1, 1, 26.86)# Name and Surname
sheet1.set_column(2, 2, 14.57)# Date
sheet1.set_column(3, 3, 18.14)# Day
sheet1.set_column(4, 4, 14.43)# Total O/T x 1  (PPH)
sheet1.set_column(5, 5, 12.57)# Total O/T x 1.5 
sheet1.set_column(6, 6, 13.57)# Total O/T x 2  




 # Write data headers.
sheet1.write('A3', 'Emp no.', headformat)
sheet1.write('B3', 'Name and Surname', headformat)
sheet1.write('C3', 'Date', headformat)
sheet1.write('D3', 'Day', headformat)
sheet1.write('E3', 'Total O/T x 1  (PPH)', headformat)
sheet1.write('F3', 'Total O/T x 1.5', headformat)
sheet1.write('G3', 'Total O/T x 2', headformat)



 # write emp no.
row = 3
col = 0
for index in range(7): 
     sheet1.write_string(row, 0, Num,fieldformat)
     sheet1.write_string(row, 1, Name,fieldformat)
     row += 1

#write Dates.
row = 4
for day in week:
    try:
        sheet1.write('C'+str(row),datetime.strptime(eval(day[:3].lower())[1], '%Y-%m-%d').strftime('%d-%b-%y'), fieldformat)
    except IndexError:
        sheet1.write('C'+str(row),'N/A', fieldformat)
    row += 1

#write week days.
row = 4
for day in week:
    sheet1.write('D'+str(row),day, fieldformat)
    row += 1

#write normal week hours.
row = 4
for day in week:
    sheet1.write('E'+str(row),'', fieldformat)
    if day == 'Sunday':
        if eval(day[:3].lower())[0]:
            sheet1.write('G10',overtime(day,eval(day[:3].lower())[2],eval(day[:3].lower())[3]), fieldformat)
        else:
            sheet1.write('G10','', fieldformat)
            
                
    else:
        if eval(day[:3].lower())[0]:
            sheet1.write('F'+str(row),overtime(day,eval(day[:3].lower())[2],eval(day[:3].lower())[3]), fieldformat)
            sheet1.write('G'+str(row),'', fieldformat)
        else:
            sheet1.write('F'+str(row),'', fieldformat)
            sheet1.write('F10','', fieldformat)
            sheet1.write('G'+str(row),'', fieldformat)
    row += 1

#write signature fields

sheet1.write('B13',Name, workbook.add_format({'bottom': 1}))
sheet1.write('B14','Prepared by', workbook.add_format({'bold': True}))
sheet1.write('B19','Manager Signature', workbook.add_format({'bold': True, 'top' : 1}))
sheet1.write('D13',str(datetime.now().strftime("%d-%b-%y")), workbook.add_format({'bottom': 1}))
sheet1.write('D14','Date', workbook.add_format({'bold': True}))
sheet1.write('D19','Date', workbook.add_format({'bold': True, 'top' : 1}))

# Notes
notesformat = workbook.add_format()
notesformat.set_center_across()
notesformat.set_bold()
notesformat.set_font_size(16)
notesformat.set_underline()

notesformat2 = workbook.add_format()
notesformat2.set_align('center_across')
#notesformat2.set_center_across()
notesformat2.set_bold()
notesformat2.set_font_size(16)

sheet1.write('D21','NOTES:',notesformat)
sheet1.write('D22','When submitting Overtime schedules ALL Time and Attendance Reports from', notesformat2)
sheet1.write('D23','Timekeeper System must be attached.', notesformat2)
sheet1.write('D24','ALL Overtime Schedules to reach Payroll Dept  on a Tuesday by 12 noon.', notesformat2)
sheet1.write('D25','All Overtime Sheets must be signed off by the respective Manager.', notesformat2)
sheet1.write('D26','PUBLIC HOLIDAYS:',notesformat)
sheet1.write('D27','Weekdays is paid at Normal Rate', notesformat2)
sheet1.write('D28','Saturdays is paid at Double Time', notesformat2)

######T&A Sheet#######
sheet2.set_column(0, 0, 50.14)# Emp info.
sheet2.set_column(1, 1, 14.86)# date
sheet2.set_column(2, 2, 8.43)# time in
sheet2.set_column(3, 3, 8.43)# time out
sheet2.set_column(4, 4, 31.71)# reason

sheet2.write('A1', 'Department Access Pattern Time Report', borbol)
sheet2.write('A2', '', borbol)

#sheet2.write('A3', 'Start Date: '+mon[1], borbol)
sheet2.write('A3', 'Start Date: '+start, borbol)
#sheet2.write('A4', 'End Date: '+sun[1], borbol)
sheet2.write('A4', 'End Date: '+str(datetime.strptime(end, '%d%m%y').strftime('%Y-%m-%d')), borbol)
sheet2.write('A5', 'Selected Person: '+Name+' ('+Num+')', borbol)
sheet2.write('A6', 'Selected Department: 51 909 2002 Information Services', borbol)
sheet2.write('A7', '', borbol)
sheet2.write('A8', 'Name', borbol)
sheet2.write('A9', 'Department: 51 909 2002 Information Services', borbol)
sheet2.write('A10', 'Name: '+Name+' ('+Num+')', borbol)

sheet2.write('B8', 'Date', borbol)
sheet2.write('C8', 'Time In', borbol)
sheet2.write('D8', 'Time Out', borbol)
sheet2.write('E8', 'Reason', borbol)

sheet2.write('B9', '', borbol)
sheet2.write('C9', '', borbol)
sheet2.write('D9', '', borbol)
sheet2.write('E9', '', borbol)

#write Dates.
row = 10
for day in week:
    try:
        sheet2.write('B'+str(row),day[:3]+' '+eval(day[:3].lower())[1], borbol)
    except IndexError:
        sheet2.write('B'+str(row),'N/A', borbol)
    row += 1



#write Times and reason.
row = 10
for day in week:
    if eval(day[:3].lower())[0]:
        sheet2.write('C'+str(row),eval(day[:3].lower())[2], workbook.add_format({'bold':True , 'border': 1,'align':'center'}))
        sheet2.write('D'+str(row),eval(day[:3].lower())[3], workbook.add_format({'bold':True , 'border': 1,'align':'center'}))
        sheet2.write('E'+str(row),'OVERTIME WORKED',workbook.add_format({ 'border': 1,'align':'center'}))
    else:
        sheet2.write('C'+str(row),'N/A', workbook.add_format({'bold':True , 'border': 1,'align':'center'}))
        sheet2.write('D'+str(row),'N/A', workbook.add_format({'bold':True , 'border': 1,'align':'center'}))
        sheet2.write('E'+str(row),'NO OVERTIME',workbook.add_format({ 'border': 1,'align':'center'}))
    row += 1




workbook.close()


##### create word .docx document

document = Document()
document.add_heading('OVERTIME SHEET', 0)
p = document.add_paragraph('Staff Name:    ')
p.add_run(Name).bold = True
p.alignment
#p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph('Week ending Sunday:    ')
p.add_run(datetime.strptime(sun[1], '%Y-%m-%d').strftime('%d %B %Y')).bold = True
#p.alignment=WD_ALIGN_PARAGRAPH.CENTER


table = document.add_table(rows=1, cols=2)
table.style = 'TableGrid'
table.cell(0,0).width = 1097280
table.cell(0,1).width = 4846320 
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('DAY').bold = True
hdr_cells[1].paragraphs[0].add_run('OVERTIME DESCRIPTION (SEE GUIDELINE BELOW)').bold = True

for day in week:
    cells = table.add_row().cells
    cells[0].text = day
    if not eval(day[:3].lower())[0]:
        cells[1].text = "NO OVERTIME"
    else:
        cells[1].text = "OVERTIME WORKED"

document.add_paragraph('Guideline entries  Below (in most cases) will apply/should suffice Use in Overtime Description', style='IntenseQuote')


document.add_paragraph('SOD', style='ListBullet')
document.add_paragraph('EOD', style='ListBullet')
document.add_paragraph('SOD/EOD', style='ListBullet')
document.add_paragraph('EOM', style='ListBullet')
p = document.add_paragraph('ERP Admin', style='ListBullet')
p.add_run('\t\t\t\t[Add brief description where possible]').bold = True
document.add_paragraph('Standby Person', style='ListBullet')
document.add_paragraph('Standby Support', style='ListBullet')
p = document.add_paragraph('Architecture Support', style='ListBullet')
p.add_run('\t\t[Add brief description where possible]').bold = True
p = document.add_paragraph('MRS Stores Support', style='ListBullet')
p.add_run('\t\t[Add brief description where possible]').bold = True
document.add_paragraph('Note:  If above does not suffice, then add your own description.', style='ListBullet')

document.add_paragraph('SPECIAL COMMENTS IF APPLICALBE/NEED MENTION STATE BELOW', style='IntenseQuote')

document.save('OT-'+Name+'-'+str(datetime.now().strftime("%d-%m-%Y"))+'.docx')




    
    
    



    
